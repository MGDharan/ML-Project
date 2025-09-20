import streamlit as st
import os
import pdfplumber
from transformers import pipeline
from docx import Document
from pdf2image import convert_from_path
import pytesseract
import tempfile
import csv
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from nltk.tokenize import sent_tokenize
import nltk

# Ensure NLTK punkt tokenizer is available
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# --- Configuration ---
st.set_page_config(page_title="AI Document Processor", layout="wide")
POPPLER_PATH = r"C:\poppler\bin"  # Change if Poppler is elsewhere
HISTORY_CSV = os.path.join(os.path.dirname(__file__), "user_history.csv")

# --- Initialize history CSV ---
if not os.path.exists(HISTORY_CSV):
    with open(HISTORY_CSV, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Timestamp", "Filename", "Operation", "Result"])

# --- Helper functions ---
def save_history(filename, operation, result):
    with open(HISTORY_CSV, "a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([datetime.now().strftime("%Y-%m-%d %H:%M:%S"), filename, operation, result])

def load_history():
    if not os.path.exists(HISTORY_CSV):
        return []
    with open(HISTORY_CSV, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        next(reader)
        return list(reader)

def extract_text_from_pdf(uploaded_file, use_ocr=False):
    text = ""
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    except:
        text = ""
    if not text.strip() and use_ocr:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded_file.read())
            tmp_path = tmp.name
        images = convert_from_path(tmp_path, poppler_path=POPPLER_PATH)
        for img in images:
            text += pytesseract.image_to_string(img)
        os.remove(tmp_path)
    return text.strip()

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def summarize_text_pages(text, total_pages):
    sentences = sent_tokenize(text)
    num_sentences = len(sentences)
    sentences_per_page = max(1, num_sentences // total_pages)
    summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
    summaries = []
    for i in range(total_pages):
        start_idx = i * sentences_per_page
        end_idx = (i + 1) * sentences_per_page if i < total_pages - 1 else num_sentences
        chunk = " ".join(sentences[start_idx:end_idx])
        try:
            summary = summarizer(chunk, max_length=300, min_length=150, do_sample=False)[0]['summary_text']
            summaries.append(summary)
        except Exception as e:
            summaries.append(f"[Error in page {i+1}: {e}]")
    return summaries

def analyze_sentiment(text):
    sentiment_pipeline = pipeline("sentiment-analysis")
    return sentiment_pipeline(text[:512])[0]

def save_output(content, filetype):
    if filetype == "PDF":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            c = canvas.Canvas(tmp.name, pagesize=letter)
            width, height = letter
            text_obj = c.beginText(40, height - 50)
            text_obj.setFont("Times-Roman", 12)
            lines = content.split("\n")
            for line in lines:
                if text_obj.getY() < 50:
                    c.drawText(text_obj)
                    c.showPage()
                    text_obj = c.beginText(40, height - 50)
                    text_obj.setFont("Times-Roman", 12)
                text_obj.textLine(line)
            c.drawText(text_obj)
            c.save()
            return tmp.name
    elif filetype == "TXT":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as tmp:
            tmp.write(content)
            return tmp.name
    elif filetype == "DOCX":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc = Document()
            doc.add_heading("AI Result", 0)
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(tmp.name)
            return tmp.name
    elif filetype == "HTML":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".html", mode="w", encoding="utf-8") as tmp:
            html_content = "<html><body><h2>AI Result</h2><p>" + content.replace("\n", "<br>") + "</p></body></html>"
            tmp.write(html_content)
            return tmp.name
    return None

# --- Sidebar ---
st.sidebar.header("Options")
uploaded_file = st.sidebar.file_uploader("Upload PDF or DOCX", type=["pdf", "docx"])
use_ocr = st.sidebar.checkbox("Use OCR (for scanned PDFs)")
operation = st.sidebar.selectbox("Operation", ["Summarize", "Sentiment Analysis", "Chatbot"])
output_format = st.sidebar.selectbox("Output format", ["PDF", "TXT", "DOCX", "HTML"])
requested_pages = st.sidebar.number_input(
    "Number of pages for summary (only for Summarize)", min_value=1, step=1, value=1
)

# --- History Sidebar ---
st.sidebar.markdown("---")
st.sidebar.subheader("History")
history = load_history()
if history:
    for entry in reversed(history[-10:]):
        st.sidebar.write(f"{entry[0]} | {entry[1]} | {entry[2]}")
else:
    st.sidebar.write("No history yet.")

# --- Main UI ---
st.title("📄 AI Document Processor")
if uploaded_file:
    filename = uploaded_file.name
    st.subheader(f"Uploaded File: {filename}")

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(uploaded_file, use_ocr=use_ocr)
        try:
            with pdfplumber.open(uploaded_file) as pdf:
                actual_pages = len(pdf.pages)
        except:
            actual_pages = max(1, len(text)//1000)
    else:
        text = extract_text_from_docx(uploaded_file)
        actual_pages = max(1, text.count("\n") // 50)

    if not text:
        st.error("❌ Could not extract text from the document.")
    else:
        st.text_area("Extracted Text Preview", text[:2000] + ("..." if len(text) > 2000 else ""), height=200)
        result = ""

        if operation == "Summarize":
            if requested_pages > actual_pages:
                st.error(f"❌ You requested {requested_pages} pages but the document has only {actual_pages} pages.")
            else:
                summaries = summarize_text_pages(text, requested_pages)
                result = "\n\n".join([f"--- Page {i+1} ---\n{s}" for i, s in enumerate(summaries)])
                st.subheader(f"📌 Summary Preview ({requested_pages} pages)")
                st.text_area("Page-wise Summary Preview", result, height=500)

        elif operation == "Sentiment Analysis":
            sentiment_result = analyze_sentiment(text)
            result = f"Label: {sentiment_result['label']}, Score: {sentiment_result['score']:.2f}"
            st.subheader("📌 Sentiment Preview")
            st.text_area("Preview Sentiment", result, height=100)

        elif operation == "Chatbot":
            user_input = st.text_input("Ask a question about the document:")
            if user_input:
                query_words = user_input.lower().split()
                doc_text_lower = text.lower()
                # Check if all keywords appear in the document
                if all(word in doc_text_lower for word in query_words):
                    result = "✅ Yes, that info is in the document."
                else:
                    result = "❌ Info not found in document."
                st.subheader("📌 Chatbot Response Preview")
                st.text_area("Preview Response", result, height=100)

        # --- Download button ---
        if operation != "Chatbot" and result:
            file_path = save_output(result, output_format)
            with open(file_path, "rb") as f:
                st.download_button(
                    f"Download Result as {output_format}",
                    data=f,
                    file_name=f"result.{output_format.lower()}"
                )
            os.remove(file_path)

        # --- Save history ---
        save_history(filename, operation, f"{requested_pages}-page summary" if operation=="Summarize" else result[:100])
