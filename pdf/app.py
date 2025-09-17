from flask import Flask, render_template, request, send_file
import os
import pdfplumber
from transformers import pipeline
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document  # for DOCX
from pdf2image import convert_from_path
import pytesseract

app = Flask(__name__)

print("1) Text Summarization\n2) Sentiment Analysis\n3) Chatbot Integration\n4) Document Classification")

# -------------------------------
# Extract text from PDF
# -------------------------------
def extract_text_from_pdf(pdf_file):
    all_text = ""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    all_text += text + "\n"
    except Exception as e:
        print(f"[Error using pdfplumber: {e}]")

    # If no text found → try OCR
    if not all_text.strip():
        print("[INFO] No selectable text found. Running OCR...")
        # Set poppler_path to uploads/poppler/bin
        poppler_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads", "poppler", "bin")
        images = convert_from_path(pdf_file, poppler_path=poppler_path)
        for img in images:
            text = pytesseract.image_to_string(img, lang="eng")
            all_text += text + "\n"

    return all_text.strip()

# -------------------------------
# Split text into chunks
# -------------------------------
def split_text(text, max_words=400):
    words = text.split()
    return [" ".join(words[i:i + max_words]) for i in range(0, len(words), max_words)]

# -------------------------------
# Summarize text (FAST MODEL)
# -------------------------------
def summarize_text(chunks):
    summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")  # Faster
    summaries = []
    for chunk in chunks:
        try:
            summary = summarizer(chunk, max_length=150, min_length=40, do_sample=False)[0]['summary_text']
            summaries.append(summary)
        except Exception as e:
            summaries.append(f"[Error summarizing chunk: {e}]")
    return "\n".join(summaries)

# -------------------------------
# Translation (placeholder)
# -------------------------------
def translate_to_language(text, language):
    if language.lower() == "tamil":
        return "இது ஒரு தமிழ் சுருக்கம் (This is a Tamil summary)."
    else:
        return f"[Translated to {language}]: {text}"

# -------------------------------
# Save summary to PDF
# -------------------------------
def save_summary_to_pdf(summary, filename="summary.pdf"):
    c = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    text_object = c.beginText(40, height - 50)
    text_object.setFont("Times-Roman", 12)

    for line in summary.split("\n"):
        text_object.textLine(line)
    c.drawText(text_object)
    c.save()
    return filename

# -------------------------------
# Save summary to TXT
# -------------------------------
def save_summary_to_txt(summary, filename="summary.txt"):
    with open(filename, "w", encoding="utf-8") as f:
        f.write(summary)
    return filename

# -------------------------------
# Save summary to DOCX
# -------------------------------
def save_summary_to_docx(summary, filename="summary.docx"):
    doc = Document()
    doc.add_heading("Summary", 0)
    for line in summary.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
    return filename

# -------------------------------
# Save summary to HTML
# -------------------------------
def save_summary_to_html(summary, filename="summary.html"):
    summary_html = summary.replace("\n", "<br>")
    html_content = f"<html><body><h2>Summary</h2><p>{summary_html}</p></body></html>"
    with open(filename, "w", encoding="utf-8") as f:
        f.write(html_content)
    return filename

# -------------------------------
# Save summary to CSV
# -------------------------------
def save_summary_to_csv(summary, filename="summary.csv"):
    with open(filename, "w", encoding="utf-8") as f:
        f.write("Summary\n")
        for line in summary.split("\n"):
            f.write(f"\"{line}\"\n")
    return filename

# -------------------------------
# Routes
# -------------------------------
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        pdf_file = request.files["pdf"]
        lang = request.form.get("language", "").strip()
        output_type = request.form.get("output_type")
        ai_tools = request.form.getlist("ai_tools")

        if pdf_file:
            upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
            os.makedirs(upload_dir, exist_ok=True)
            filepath = os.path.join(upload_dir, pdf_file.filename)
            pdf_file.save(filepath)

            # Extract text
            text = extract_text_from_pdf(filepath)
            print(f"[DEBUG] Extracted text length: {len(text)}")
            if not text.strip():
                return render_template("index.html", error="❌ Could not extract text from PDF (even with OCR).", ai_result=None)

            result_sections = []

            # 1. Text Summarization
            if "summarization" in ai_tools:
                chunks = split_text(text, max_words=100)
                summary = summarize_text(chunks)
                result_sections.append("--- Text Summarization ---\n" + summary)
            else:
                summary = text

            # 2. Sentiment Analysis
            if "sentiment" in ai_tools:
                sentiment_analyzer = pipeline("sentiment-analysis")
                sentiment = sentiment_analyzer(summary)[0]
                sentiment_str = f"Label: {sentiment['label']}, Score: {sentiment['score']:.2f}"
                result_sections.append("--- Sentiment Analysis ---\n" + sentiment_str)

            # 3. Chatbot Integration (simple echo)
            if "chatbot" in ai_tools:
                chatbot_response = f"Chatbot: You said '{summary[:100]}...'"
                result_sections.append("--- Chatbot Integration ---\n" + chatbot_response)

            # 4. Document Classification
            if "classification" in ai_tools:
                classifier = pipeline("zero-shot-classification")
                candidate_labels = ["business", "education", "health", "technology", "finance", "other"]
                classification = classifier(summary, candidate_labels)
                top_label = classification['labels'][0]
                top_score = classification['scores'][0]
                class_str = f"Top class: {top_label} (score: {top_score:.2f})"
                result_sections.append("--- Document Classification ---\n" + class_str)

            # If no AI tool selected, just use summary
            if not result_sections:
                result_sections.append(summary)

            final_output = "\n\n".join(result_sections)

            # Translate (optional)
            if lang:
                final_output = translate_to_language(final_output, lang)

            # Save based on format (absolute paths)
            if output_type == "pdf":
                out_file = os.path.join(upload_dir, "final_summary.pdf")
                save_summary_to_pdf(final_output, out_file)
            elif output_type == "txt":
                out_file = os.path.join(upload_dir, "final_summary.txt")
                save_summary_to_txt(final_output, out_file)
            elif output_type == "doc":
                out_file = os.path.join(upload_dir, "final_summary.docx")
                save_summary_to_docx(final_output, out_file)
            elif output_type == "html":
                out_file = os.path.join(upload_dir, "final_summary.html")
                save_summary_to_html(final_output, out_file)
            elif output_type == "csv":
                out_file = os.path.join(upload_dir, "final_summary.csv")
                save_summary_to_csv(final_output, out_file)
            else:
                out_file = os.path.join(upload_dir, "final_summary.txt")
                save_summary_to_txt(final_output, out_file)

            # Error handling for missing file
            if not os.path.exists(out_file):
                return render_template("index.html", error="Error: Output file was not created.", ai_result=None)
            print(f"[DEBUG] Saved file: {out_file}")
            # Show result and download link
            download_link = f"/download/{os.path.basename(out_file)}"
            return render_template("index.html", ai_result=final_output, download_link=download_link, error=None)

    return render_template("index.html", ai_result=None, download_link=None, error=None)

@app.route("/download/<filename>")
def download_file(filename):
    upload_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads")
    file_path = os.path.join(upload_dir, filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    return "File not found", 404

if __name__ == "__main__":
    app.run(debug=True)
